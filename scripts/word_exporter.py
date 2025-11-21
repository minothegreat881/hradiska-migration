"""
Word Exporter pre hradiska.sk
Export všetkého obsahu do Word dokumentov
"""

import json
import os
from pathlib import Path
from datetime import datetime
from typing import List, Dict
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_HEADING_LEVEL
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import requests
from PIL import Image
import io
import re
import markdown
from bs4 import BeautifulSoup

class WordExporter:
    def __init__(self, content_dir: str = "../nextjs-app/content", output_dir: str = "../word-export"):
        self.content_dir = Path(content_dir)
        self.output_dir = Path(output_dir)
        self.articles_dir = self.output_dir / "articles"
        self.complete_dir = self.output_dir / "complete"
        self.images_dir = self.output_dir / "images"

        # Vytvorenie adresárov
        for dir_path in [self.articles_dir, self.complete_dir, self.images_dir]:
            dir_path.mkdir(parents=True, exist_ok=True)

        # Štýly dokumentu
        self.styles = {
            'primary_color': RGBColor(139, 69, 19),  # Hnedá
            'secondary_color': RGBColor(34, 139, 34),  # Zelená
            'text_color': RGBColor(51, 51, 51),  # Tmavošedá
            'heading_font': 'Georgia',
            'body_font': 'Calibri'
        }

    def setup_document_styles(self, doc: Document):
        """Nastaví štýly dokumentu"""
        # Štýl pre nadpis 1
        heading1_style = doc.styles['Heading 1']
        heading1_style.font.name = self.styles['heading_font']
        heading1_style.font.size = Pt(24)
        heading1_style.font.color.rgb = self.styles['primary_color']
        heading1_style.font.bold = True

        # Štýl pre nadpis 2
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = self.styles['heading_font']
        heading2_style.font.size = Pt(18)
        heading2_style.font.color.rgb = self.styles['primary_color']

        # Štýl pre nadpis 3
        heading3_style = doc.styles['Heading 3']
        heading3_style.font.name = self.styles['heading_font']
        heading3_style.font.size = Pt(14)
        heading3_style.font.color.rgb = self.styles['secondary_color']

        # Štýl pre normálny text
        normal_style = doc.styles['Normal']
        normal_style.font.name = self.styles['body_font']
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = self.styles['text_color']

    def add_title_page(self, doc: Document):
        """Pridá titulnú stranu"""
        # Hlavný nadpis
        title = doc.add_heading('Slovanské Hradiská', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Podnadpis
        subtitle = doc.add_paragraph('Kompletná dokumentácia historických hradísk')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].font.italic = True

        doc.add_paragraph()
        doc.add_paragraph()

        # Popis
        description = doc.add_paragraph(
            'Tento dokument obsahuje kompletnú dokumentáciu webstránky hradiska.sk, '
            'vrátane všetkých článkov, obrázkov a historických informácií o slovanských '
            'hradiskách na území Slovenska.'
        )
        description.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph()
        doc.add_paragraph()

        # Informácie o exporte
        info = doc.add_paragraph(f'Dátum exportu: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info.runs[0].font.size = Pt(10)

        # Nová strana
        doc.add_page_break()

    def add_table_of_contents(self, doc: Document, articles: List[Dict]):
        """Pridá obsah"""
        doc.add_heading('Obsah', 1)

        # Hlavné sekcie
        sections = {
            'historia': 'Historické články',
            'archeologia': 'Archeologické nálezy',
            'mytologia': 'Mytológia a kultúra',
            'mapa': 'Mapy a lokality',
            'ostatne': 'Ostatné články'
        }

        for section_id, section_title in sections.items():
            # Nadpis sekcie
            doc.add_heading(section_title, 2)

            # Články v sekcii
            section_articles = [a for a in articles if self.categorize_article(a) == section_id]
            for article in section_articles:
                p = doc.add_paragraph()
                p.add_run(f"• {article['title']}")
                p.paragraph_format.left_indent = Inches(0.5)

        doc.add_page_break()

    def categorize_article(self, article: Dict) -> str:
        """Kategorizuje článok na základe obsahu"""
        title = article.get('title', '').lower()
        categories = [cat.lower() for cat in article.get('categories', [])]
        tags = [tag.lower() for tag in article.get('tags', [])]

        all_text = ' '.join([title] + categories + tags)

        if any(word in all_text for word in ['história', 'historický', 'dejiny', 'veľká morava']):
            return 'historia'
        elif any(word in all_text for word in ['archeologický', 'nález', 'vykopávky', 'hrob']):
            return 'archeologia'
        elif any(word in all_text for word in ['mytológia', 'bohovia', 'mýtus', 'povesť']):
            return 'mytologia'
        elif any(word in all_text for word in ['mapa', 'lokalita', 'miesto', 'poloha']):
            return 'mapa'
        else:
            return 'ostatne'

    def markdown_to_word(self, md_text: str, doc: Document):
        """Konvertuje Markdown text do Word dokumentu"""
        # Konverzia Markdown na HTML
        html = markdown.markdown(md_text, extensions=['extra', 'codehilite', 'tables'])
        soup = BeautifulSoup(html, 'html.parser')

        for element in soup.find_all():
            if element.name == 'h1':
                doc.add_heading(element.get_text(), 1)
            elif element.name == 'h2':
                doc.add_heading(element.get_text(), 2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), 3)
            elif element.name == 'p':
                paragraph = doc.add_paragraph()
                for child in element.children:
                    if child.name == 'strong' or child.name == 'b':
                        paragraph.add_run(child.get_text()).bold = True
                    elif child.name == 'em' or child.name == 'i':
                        paragraph.add_run(child.get_text()).italic = True
                    elif child.name == 'a':
                        run = paragraph.add_run(child.get_text())
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        run.underline = True
                    else:
                        paragraph.add_run(str(child))
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(li.get_text())
            elif element.name == 'ol':
                for i, li in enumerate(element.find_all('li'), 1):
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(li.get_text())
            elif element.name == 'blockquote':
                p = doc.add_paragraph()
                p.add_run(element.get_text()).italic = True
                p.paragraph_format.left_indent = Inches(0.5)
            elif element.name == 'table':
                self.add_table_from_html(doc, element)

    def add_table_from_html(self, doc: Document, table_element):
        """Pridá tabuľku z HTML do dokumentu"""
        rows = table_element.find_all('tr')
        if not rows:
            return

        # Počet stĺpcov
        cols = max(len(row.find_all(['td', 'th'])) for row in rows)

        # Vytvorenie tabuľky
        table = doc.add_table(rows=0, cols=cols)
        table.style = 'Light Grid Accent 1'

        # Pridanie riadkov
        for row_element in rows:
            row = table.add_row()
            cells = row_element.find_all(['td', 'th'])
            for i, cell in enumerate(cells):
                if i < cols:
                    row.cells[i].text = cell.get_text().strip()

    def add_image(self, doc: Document, image_path: str, caption: str = None, max_width: float = 6.0):
        """Pridá obrázok do dokumentu"""
        try:
            # Otvorenie obrázka
            img = Image.open(image_path)

            # Zmena veľkosti ak je príliš veľký
            width_inches = img.width / 96  # 96 DPI
            if width_inches > max_width:
                doc.add_picture(image_path, width=Inches(max_width))
            else:
                doc.add_picture(image_path)

            # Pridanie popisu
            if caption:
                p = doc.add_paragraph()
                p.add_run(f"Obrázok: {caption}").italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        except Exception as e:
            print(f"Chyba pri pridávaní obrázka {image_path}: {e}")

    def export_single_article(self, article_data: Dict, output_path: Path):
        """Exportuje jednotlivý článok do Word dokumentu"""
        doc = Document()
        self.setup_document_styles(doc)

        # Nadpis
        doc.add_heading(article_data.get('title', 'Bez názvu'), 1)

        # Metadáta
        metadata = doc.add_paragraph()
        metadata.add_run(f"Dátum: {article_data.get('date', 'Neznámy')}\n")
        metadata.add_run(f"Autor: {article_data.get('author', 'hradiska.sk')}\n")
        if article_data.get('categories'):
            metadata.add_run(f"Kategórie: {', '.join(article_data['categories'])}\n")
        if article_data.get('tags'):
            metadata.add_run(f"Tagy: {', '.join(article_data['tags'])}")
        metadata.runs[0].font.size = Pt(10)
        metadata.runs[0].font.italic = True

        doc.add_paragraph()

        # Obsah
        content = article_data.get('content', '')
        if content:
            self.markdown_to_word(content, doc)

        # Obrázky
        if article_data.get('images'):
            doc.add_heading('Obrázky', 2)
            for img_data in article_data['images']:
                # Tu by sme mali stiahnuť a pridať obrázky
                doc.add_paragraph(f"[Obrázok: {img_data.get('alt', img_data.get('src', ''))}]")

        # Uloženie
        doc.save(output_path)

    def export_all_articles(self):
        """Exportuje všetky články do samostatných Word dokumentov"""
        # Načítanie zoznamu článkov
        articles_file = self.content_dir / "data" / "articles.json"
        if not articles_file.exists():
            print("Súbor s článkami neexistuje!")
            return

        with open(articles_file, 'r', encoding='utf-8') as f:
            articles = json.load(f)

        print(f"Exportujem {len(articles)} článkov...")

        for i, article_info in enumerate(articles, 1):
            print(f"Exportujem {i}/{len(articles)}: {article_info['title']}")

            # Načítanie MDX súboru
            mdx_file = self.content_dir / "posts" / article_info['file']
            if not mdx_file.exists():
                print(f"Súbor {mdx_file} neexistuje!")
                continue

            with open(mdx_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # Rozdelenie na frontmatter a obsah
            if content.startswith('---'):
                parts = content.split('---', 2)
                if len(parts) >= 3:
                    article_content = parts[2]
                else:
                    article_content = content
            else:
                article_content = content

            # Príprava dát článku
            article_data = {
                'title': article_info['title'],
                'date': article_info.get('date', ''),
                'categories': article_info.get('categories', []),
                'tags': article_info.get('tags', []),
                'content': article_content
            }

            # Export do Word
            output_file = self.articles_dir / f"{article_info['slug']}.docx"
            self.export_single_article(article_data, output_file)

        print("Export jednotlivých článkov dokončený!")

    def export_complete_document(self):
        """Vytvorí jeden kompletný Word dokument so všetkým obsahom"""
        print("Vytváram kompletný dokument...")

        doc = Document()
        self.setup_document_styles(doc)

        # Titulná strana
        self.add_title_page(doc)

        # Načítanie článkov
        articles_file = self.content_dir / "data" / "articles.json"
        if not articles_file.exists():
            print("Súbor s článkami neexistuje!")
            return

        with open(articles_file, 'r', encoding='utf-8') as f:
            articles = json.load(f)

        # Obsah
        self.add_table_of_contents(doc, articles)

        # Sekcie
        sections = {
            'historia': 'Historické články',
            'archeologia': 'Archeologické nálezy',
            'mytologia': 'Mytológia a kultúra',
            'mapa': 'Mapy a lokality',
            'ostatne': 'Ostatné články'
        }

        for section_id, section_title in sections.items():
            # Nadpis sekcie
            doc.add_heading(section_title, 1)

            # Články v sekcii
            section_articles = [a for a in articles if self.categorize_article(a) == section_id]

            for article_info in section_articles:
                # Nadpis článku
                doc.add_heading(article_info['title'], 2)

                # Načítanie obsahu
                mdx_file = self.content_dir / "posts" / article_info['file']
                if mdx_file.exists():
                    with open(mdx_file, 'r', encoding='utf-8') as f:
                        content = f.read()

                    # Rozdelenie na frontmatter a obsah
                    if content.startswith('---'):
                        parts = content.split('---', 2)
                        if len(parts) >= 3:
                            article_content = parts[2]
                        else:
                            article_content = content
                    else:
                        article_content = content

                    # Pridanie obsahu
                    self.markdown_to_word(article_content, doc)

                doc.add_page_break()

        # Záverečná strana
        doc.add_heading('Záver', 1)
        doc.add_paragraph(
            'Tento dokument obsahuje kompletnú dokumentáciu webstránky hradiska.sk. '
            'Všetky informácie boli exportované z originálnej webstránky a prevedené '
            'do formátu Microsoft Word pre archiváciu a ďalšie spracovanie.'
        )

        # Uloženie
        output_file = self.complete_dir / f"hradiska_kompletny_dokument_{datetime.now().strftime('%Y%m%d')}.docx"
        doc.save(output_file)

        print(f"Kompletný dokument uložený: {output_file}")

    def create_index_document(self):
        """Vytvorí index dokument so zoznamom všetkých súborov"""
        doc = Document()
        self.setup_document_styles(doc)

        doc.add_heading('Index exportovaných dokumentov', 1)

        # Zoznam jednotlivých článkov
        doc.add_heading('Jednotlivé články', 2)
        for docx_file in sorted(self.articles_dir.glob('*.docx')):
            p = doc.add_paragraph()
            p.add_run(f"• {docx_file.stem}")

        # Informácie o exporte
        doc.add_heading('Informácie o exporte', 2)
        info_para = doc.add_paragraph()
        info_para.add_run(f"Dátum exportu: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
        info_para.add_run(f"Počet článkov: {len(list(self.articles_dir.glob('*.docx')))}\n")
        info_para.add_run(f"Umiestnenie: {self.output_dir.absolute()}")

        # Uloženie
        index_file = self.output_dir / "INDEX.docx"
        doc.save(index_file)
        print(f"Index dokument vytvorený: {index_file}")

def main():
    """Hlavná funkcia"""
    exporter = WordExporter()

    print("Začínam export do Word dokumentov...")
    print("=" * 50)

    # Export jednotlivých článkov
    print("\n1. Export jednotlivých článkov...")
    exporter.export_all_articles()

    # Vytvorenie kompletného dokumentu
    print("\n2. Vytváram kompletný dokument...")
    exporter.export_complete_document()

    # Vytvorenie indexu
    print("\n3. Vytváram index dokument...")
    exporter.create_index_document()

    print("\n" + "=" * 50)
    print("Export dokončený!")
    print(f"Dokumenty sú uložené v: {exporter.output_dir}")

if __name__ == "__main__":
    main()